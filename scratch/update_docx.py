import docx

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = docx.oxml.parse_xml(f'<w:shd {docx.oxml.ns.nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_table_spec(doc, title, fields):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = docx.shared.Pt(12)
    
    table = doc.add_table(rows=1, cols=5)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ['Field', 'Field2 (한글명)', 'Domain', 'Type', 'Key']
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_background(hdr_cells[idx], 'F2F2F2')
    
    for row_data in fields:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row_data):
            row_cells[idx].text = val
    doc.add_paragraph('')

src_path = 'c:/Runners_plan/산출물/테이블_명세서_Running_Coach_v2.0.docx'
dst_path = 'c:/Runners_plan/산출물/PostgreSQL_테이블_명세서_Running_Coach_v2.0.docx'

doc = docx.Document(src_path)

garmin_detail_fields = [
    ['garmin_detail_id', '가민상세ID', 'PK', 'BIGINT', 'PK'],
    ['run_record_id', '러닝기록ID', 'FK', 'BIGINT', 'FK (UQ)'],
    ['garmin_activity_id', '가민활동ID', 'ID', 'BIGINT', 'UQ'],
    ['max_hr', '최대심박수', 'HEART_RATE', 'INTEGER', ''],
    ['avg_cadence', '평균케이던스', 'CADENCE', 'INTEGER', ''],
    ['max_cadence', '최대케이던스', 'CADENCE', 'INTEGER', ''],
    ['avg_stride_length_mm', '평균보폭(mm)', 'LENGTH', 'INTEGER', ''],
    ['elevation_gain_m', '획득고도(m)', 'ELEVATION', 'NUMERIC(6,1)', ''],
    ['elevation_loss_m', '손실고도(m)', 'ELEVATION', 'NUMERIC(6,1)', ''],
    ['vo2_max', '추정VO2Max', 'SCORE', 'NUMERIC(4,1)', ''],
    ['training_effect_aerobic', '유산소훈련효과', 'EFFECT', 'NUMERIC(3,1)', ''],
    ['training_effect_anaerobic', '무산소훈련효과', 'EFFECT', 'NUMERIC(3,1)', ''],
    ['calories', '소모칼로리(kcal)', 'CALORIE', 'INTEGER', ''],
    ['gpx_route_json', 'GPS경로JSON', 'JSON', 'JSONB', ''],
    ['created_at', '생성일시', 'DATETIME', 'TIMESTAMP', ''],
    ['updated_at', '수정일시', 'DATETIME', 'TIMESTAMP', '']
]

garmin_lap_fields = [
    ['lap_id', '랩ID', 'PK', 'BIGINT', 'PK'],
    ['run_record_id', '러닝기록ID', 'FK', 'BIGINT', 'FK'],
    ['lap_index', '구간순번', 'INDEX', 'INTEGER', ''],
    ['lap_distance_km', '구간거리(km)', 'DISTANCE', 'NUMERIC(5,2)', ''],
    ['lap_duration_sec', '구간소요시간(초)', 'DURATION', 'INTEGER', ''],
    ['lap_avg_pace_sec', '구간평균페이스(초)', 'PACE', 'INTEGER', ''],
    ['lap_avg_hr', '구간평균심박수', 'HEART_RATE', 'INTEGER', ''],
    ['lap_max_hr', '구간최대심박수', 'HEART_RATE', 'INTEGER', ''],
    ['lap_avg_cadence', '구간케이던스', 'CADENCE', 'INTEGER', ''],
    ['created_at', '생성일시', 'DATETIME', 'TIMESTAMP', '']
]

add_table_spec(doc, 'GARMIN_RUN_DETAIL (가민 정밀 분석 상세)', garmin_detail_fields)
add_table_spec(doc, 'GARMIN_RUN_LAP (가민 km/구간 랩 타임)', garmin_lap_fields)

doc.save(dst_path)
print('Successfully generated full PostgreSQL_테이블_명세서_Running_Coach_v2.0.docx!')
