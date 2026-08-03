import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helpers for XML manipulations
def set_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def add_bottom_border(paragraph, color_hex="1A365D", size="12"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)  # 12 = 1.5 pt
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    # Padding in dxa (1 pt = 20 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style is not None:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), border_style.get('val', 'single'))
            border_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(border_el)
        else:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
    tcPr.append(tcBorders)

def set_col_widths(table, widths):
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Primary Navy
    
    add_bottom_border(p, color_hex="1A365D", size="8")
    return p

def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    # Background light gray-blue shading
    set_cell_shading(cell, "F7FAFC")
    
    # Left border: thick blue, others: none
    set_cell_borders(cell, 
                     left={'val': 'single', 'sz': 24, 'color': '2B6CB0'}, # 3pt
                     top=None, bottom=None, right=None)
    
    # Padding
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Set table width (160mm total width)
    cell.width = Mm(160)
    
    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)

def add_custom_bullet(doc, lead_in, description, num_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.first_line_indent = Mm(-6)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    
    # Bullet/Number prefix
    if num_label:
        prefix_run = p.add_run(f"{num_label}  ")
        prefix_run.font.bold = True
        prefix_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    else:
        prefix_run = p.add_run("•  ")
        prefix_run.font.bold = True
        prefix_run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    set_font(prefix_run, "Arial")
    
    # Lead-in (bold)
    if lead_in:
        lead_run = p.add_run(lead_in)
        set_font(lead_run, "Malgun Gothic")
        lead_run.font.bold = True
        lead_run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
    # Description (normal)
    desc_run = p.add_run(description)
    set_font(desc_run, "Malgun Gothic")
    desc_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slightly lighter slate gray

def add_formatted_text(paragraph, text):
    parts = text.split('**')
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        set_font(run, "Malgun Gothic")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        if idx % 2 == 1:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Emphasized in primary color

def create_info_table(doc, data, col_widths):
    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, (label, val) in enumerate(data):
        row_cells = table.rows[r_idx].cells
        
        # Label cell
        cell_lbl = row_cells[0]
        set_cell_shading(cell_lbl, "F0F4F8")
        set_cell_margins(cell_lbl, top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell_lbl, 
                         top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         left=None, right=None)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_after = Pt(0)
        run_lbl = p_lbl.add_run(label)
        set_font(run_lbl, "Malgun Gothic")
        run_lbl.font.size = Pt(10)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        # Value cell
        cell_val = row_cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_margins(cell_val, top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell_val, 
                         top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                         left=None, right=None)
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_val.paragraph_format.line_spacing = 1.25
        p_val.paragraph_format.space_after = Pt(0)
        run_val = p_val.add_run(val)
        set_font(run_val, "Malgun Gothic")
        run_val.font.size = Pt(10)
        run_val.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
    set_col_widths(table, col_widths)
    
    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)

def create_styled_table(doc, headers, data, col_widths):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, "1A365D")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        set_cell_borders(cell, 
                         top={'val': 'single', 'sz': 4, 'color': '1A365D'},
                         bottom={'val': 'single', 'sz': 8, 'color': '1A365D'},
                         left=None, right=None)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        set_font(run, "Malgun Gothic")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, text in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            set_cell_borders(cell, 
                             top=None,
                             bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             left=None, right=None)
            
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(0)
            
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            add_formatted_text(p, text)
            
    set_col_widths(table, col_widths)
    
    # Spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)

def main():
    doc = Document()
    
    # Configure A4 page and 25mm margins
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        
    # Set base font on normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("프로젝트 개요서")
    set_font(title_run, "Malgun Gothic")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(20)
    sub_run = subtitle_p.add_run("Running Coach - 러닝 기록 및 훈련 분석 플랫폼")
    set_font(sub_run, "Malgun Gothic")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    add_bottom_border(subtitle_p, color_hex="D2D6DC", size="6")
    
    # Section 1: 프로젝트 정보
    add_heading_1(doc, "1. 프로젝트 정보")
    info_data = [
        ("프로젝트명", "Running Coach"),
        ("프로젝트 기간", "2026.06 ~ 2026.09"),
        ("개발 인원", "1명"),
        ("개발 목적", "러닝 기록과 목표를 체계적으로 관리하고 데이터 분석을 통해 훈련 성과를 확인할 수 있는 플랫폼 개발")
    ]
    create_info_table(doc, info_data, [Mm(35), Mm(125)])
    
    # Section 2: 프로젝트 배경
    add_heading_1(doc, "2. 프로젝트 배경")
    add_callout(doc, "기존의 대다수 러닝 애플리케이션은 단순한 활동 기록 저장 중심의 기능만을 제공하고 있어, "
                     "자신의 한계에 도전하고 기량을 발전시키고자 하는 전문적/취미 러너들의 다각화된 요구사항을 "
                     "충족시키기에 부족함이 있습니다.")
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(6)
    run_intro = p_intro.add_run("러너들은 단순히 달린 거리를 기록하는 것을 넘어 다음과 같은 체계적인 기능들을 필요로 하고 있습니다:")
    set_font(run_intro, "Malgun Gothic")
    run_intro.font.size = Pt(10)
    
    add_custom_bullet(doc, "월별 목표 달성률: ", "지속적인 훈련 동기부여와 페이스 관리를 위한 월별/기간별 목표치 대비 진척도 시각화")
    add_custom_bullet(doc, "대회 기록 관리: ", "참가 예정이거나 완료된 공식 레이스 성적 및 기록증의 체계적 보관과 히스토리 조회")
    add_custom_bullet(doc, "훈련 분석: ", "단순 평균 페이스 외에 구간별 페이스 분석, 심박수 트렌드 분석 등 과학적 피드백 제공")
    
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_before = Pt(8)
    p_summary.paragraph_format.space_after = Pt(12)
    run_summary = p_summary.add_run("따라서 본 프로젝트는 러너의 실제 훈련 데이터를 체계적으로 축적하고 분석할 수 있는 맞춤형 분석 대시보드를 제공하여 "
                                     "체계적인 러닝 성장을 돕는 웹 플랫폼을 구축하고자 합니다.")
    set_font(run_summary, "Malgun Gothic")
    run_summary.font.size = Pt(10)
    
    # Section 3: 프로젝트 목표
    add_heading_1(doc, "3. 프로젝트 목표")
    goals = [
        ("러닝 기록 관리", "활동 기록의 체계적인 등록, 수정 및 세부 페이스 데이터 조회 환경 제공"),
        ("목표 관리", "월별 누적 목표 설정 및 진척도 시뮬레이션 기능 구축"),
        ("대회 기록 관리", "참가 대회 일정 관리 및 공식 기록증 이미지 업로드/디지털 보관"),
        ("통계 분석", "주간/월간 러닝 패턴 및 누적 통계 데이터를 시각화 대시보드로 전달"),
        ("개인 기록(PB) 관리", "거리별(5km, 10km, Half, Full) 개인 최고 기록 자동 판별 및 히스토리 제공")
    ]
    for idx, (title, desc) in enumerate(goals, 1):
        add_custom_bullet(doc, f"{title}: ", desc, num_label=f"{idx}.")
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Section 4: 주요 기능
    add_heading_1(doc, "4. 주요 기능")
    feature_headers = ["구분", "주요 기능 및 상세 설명"]
    feature_data = [
        ("회원", "**소셜 로그인**: 카카오, 네이버, 구글 등 소셜 간편 연동을 통한 손쉬운 접근성 보장"),
        ("러닝기록", "**기록 등록/수정/삭제**: 달린 날짜, 시간, 거리, 평균 페이스 및 개인 상태 메모 작성"),
        ("목표관리", "**월별 목표 설정**: 월간 누적 달리기 목표 거리 설정 및 진행 현황 그래프 표시"),
        ("목표관리", "**레이스 목표 설정**: 다가오는 타겟 대회 및 목표 완주 시간, 페이스 계획 수립"),
        ("대회관리", "**기록증 관리**: 참여 완료된 대회 성적 기록 및 실물 기록증 이미지 업로드/관리"),
        ("통계", "**PB 분석**: 개인 최고 기록(PB) 자동 업데이트 및 구간별 달성 트렌드 통계"),
        ("통계", "**월별 분석**: 월별 총 활동 거리, 평균 페이스, 소모 칼로리 등 다차원 시각화 분석")
    ]
    create_styled_table(doc, feature_headers, feature_data, [Mm(35), Mm(125)])
    
    # Section 5: 기대 효과
    add_heading_1(doc, "5. 기대 효과")
    effects = [
        ("러닝 데이터의 통합 관리: ", "여러 기기에 분산되어 있거나 단편적인 기록을 단일 플랫폼에서 체계적인 포맷으로 아카이빙 가능"),
        ("목표 기반의 동기부여 강화: ", "월별 목표 대비 실시간 달성률 분석을 제공하여 지속적인 훈련 참여도 고취"),
        ("대회 및 성장 히스토리 보관: ", "개인의 공식 대회 성적과 기록증을 함께 보존함으로써 개인 러닝 역사 관리 용이"),
        ("데이터 기반의 기량 향상: ", "평균 페이스 및 훈련 패턴 통계를 시각적으로 피드백받아 효율적인 페이스 전략 수립 가능")
    ]
    for title, desc in effects:
        add_custom_bullet(doc, title, desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Section 6: 기술 스택
    add_heading_1(doc, "6. 기술 스택")
    tech_headers = ["구분", "기술 스택", "상세 역할 및 선정 이유"]
    tech_data = [
        ("Backend", "Spring Boot", "안정적인 비즈니스 로직 처리 및 RESTful API 구현, 데이터 보안 관리"),
        ("Frontend", "React", "컴포넌트 단위의 유연한 화면 설계 및 반응형 UI 제공으로 최상의 대시보드 UX 구현"),
        ("Database", "PostgreSQL", "강력한 RDBMS 기능을 활용한 사용자 기록의 정합성 보장 및 통계 데이터 쿼리 최적화"),
        ("Authentication", "OAuth2", "소셜 로그인 연동(OAuth 2.0 프로토콜)을 통한 빠르고 안전한 인증 관리"),
        ("Deploy", "Docker", "개발 및 운영 환경 불일치 방지와 효율적인 컨테이너 기반 배포 인프라 제공"),
        ("Test", "JUnit, Playwright", "백엔드 API 단위 테스트 및 사용자 관점의 화면 흐름 E2E 테스트 자동화")
    ]
    create_styled_table(doc, tech_headers, tech_data, [Mm(30), Mm(40), Mm(90)])
    
    # Section 7: 향후 확장 계획
    add_heading_1(doc, "7. 향후 확장 계획")
    extensions = [
        ("FIT 파일 업로드 지원: ", "가민(Garmin) 또는 스트라바(Strava) 기기에서 다운로드받는 GPS .fit 파일 자동 업로드 및 상세 지도 경로 매핑"),
        ("자동 랩(Lap) 페이스 분석: ", "1km 단위의 자동 랩 분석 및 고저차가 페이스에 미치는 영향 자동 연산"),
        ("AI 훈련 코치 서비스: ", "사용자의 누적 페이스 데이터를 바탕으로 다음 목표 대회까지의 주차별 추천 훈련 스케줄 및 러닝 조언 생성"),
        ("날씨 데이터 자동 수집: ", "러닝 기록 일시를 기준으로 공공 기상 API에서 온도, 습도, 풍향 정보를 조회하여 기록에 자동 연동"),
        ("러닝 레벨 및 뱃지 시스템: ", "누적 거리 및 특정 챌린지 달성 시 리워드 성격의 등급 부여와 한정판 디지털 뱃지 획득을 통한 게이미케이션")
    ]
    for idx, (title, desc) in enumerate(extensions, 1):
        add_custom_bullet(doc, title, desc, num_label=f"{idx}.")
        
    # Document Footer / Info
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(36)
    run_footer = p_footer.add_run("문서 작성일: 2026. 06. 18  |  작성자: Running Coach 개발팀")
    set_font(run_footer, "Malgun Gothic")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    filename = "프로젝트_개요서_Running_Coach.docx"
    doc.save(filename)
    print(f"Successfully created '{filename}'")

if __name__ == "__main__":
    main()
