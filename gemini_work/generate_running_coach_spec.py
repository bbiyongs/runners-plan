import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helper: Set font for runs including East Asian (Korean) fonts
def set_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# Helper: Add custom bottom border to paragraphs for titles
def add_bottom_border(paragraph, color_hex="1A365D", size="12"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)  # 12 = 1.5 pt
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

# Helper: Set cell shading background color
def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

# Helper: Set custom padding for table cells
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper: Set custom cell borders
def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style is not None:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), border_style.get('val', 'single'))
            border_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(border_el)
        else:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
    tcPr.append(tcBorders)

# Helper: Set column widths for tables
def set_col_widths(table, widths):
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

# Helper: Add Heading 1 with Bottom Line
def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    add_bottom_border(p, color_hex="1A365D", size="8")
    return p

# Helper: Add Heading 2
def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    return p

# Helper: Add a custom list/bullet point
def add_custom_bullet(doc, text, num_label=None, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6 * (level + 1))
    p.paragraph_format.first_line_indent = Mm(-6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    
    if num_label:
        prefix = f"{num_label}  "
        color = RGBColor(0x1A, 0x36, 0x5D)
    else:
        prefix = "-  "
        color = RGBColor(0x2B, 0x6C, 0xB0)
        
    prefix_run = p.add_run(prefix)
    prefix_run.font.bold = True
    prefix_run.font.color.rgb = color
    set_font(prefix_run, "Arial" if not num_label else "Malgun Gothic")
    
    desc_run = p.add_run(text)
    set_font(desc_run, "Malgun Gothic")
    desc_run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

# Helper: Create standard styled 2-column table
def create_styled_table(doc, headers, data, col_widths):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, "1A365D")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        set_cell_borders(cell, 
                         top={'val': 'single', 'sz': 4, 'color': '1A365D'},
                         bottom={'val': 'single', 'sz': 8, 'color': '1A365D'},
                         left=None, right=None)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        set_font(run, "Malgun Gothic")
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, text in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_borders(cell, 
                             top=None,
                             bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             left=None, right=None)
            
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(0)
            
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_font(run, "Malgun Gothic")
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(text)
                set_font(run, "Malgun Gothic")
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                
    set_col_widths(table, col_widths)
    
    # Bottom margin spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)

def main():
    doc = Document()
    
    # Page setup - A4 with 25mm margins
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        
    # Default Paragraph Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Document Main Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(15)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("Running Coach 프로젝트 기획 및 설계안")
    set_font(title_run, "Malgun Gothic")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(20)
    sub_run = subtitle_p.add_run("요청 원본 텍스트 기반 작성 문서")
    set_font(sub_run, "Malgun Gothic")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    add_bottom_border(subtitle_p, color_hex="D2D6DC", size="6")
    
    # 1. 프로젝트명 (가칭)
    add_heading_1(doc, "프로젝트명 (가칭)")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    run = p.add_run("Running Coach")
    set_font(run, "Malgun Gothic")
    
    # 2. 프로젝트 소개
    add_heading_1(doc, "프로젝트 소개")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("러너가 일일 훈련 기록과 대회 기록을 관리하고, 월별 목표 및 레이스 목표를 설정하여 달성률을 확인할 수 있으며, 훈련 데이터를 분석하여 자신의 러닝 패턴과 성장 과정을 확인할 수 있는 플랫폼")
    set_font(run, "Malgun Gothic")
    
    # 3. 프로젝트 목표
    add_heading_1(doc, "프로젝트 목표")
    objectives = [
        "러닝 기록 관리",
        "월별 목표 및 레이스 목표 관리",
        "대회 기록 및 기록증 관리",
        "개인 베스트(PB) 관리",
        "훈련 데이터 분석 및 시각화",
        "향후 FIT 파일 연동 및 AI 러닝 코치 기능 확장"
    ]
    for idx, obj in enumerate(objectives, 1):
        add_custom_bullet(doc, obj, num_label=f"{idx}.")
        
    # 4. 기술 스택 (예정)
    add_heading_1(doc, "기술 스택 (예정)")
    tech_headers = ["구분", "기술 스택"]
    tech_data = [
        ("Backend", "Spring Boot, Spring Security, JWT, MyBatis"),
        ("Frontend", "React"),
        ("Database", "MySQL / Oracle"),
        ("API", "OAuth2 (Google / Naver)"),
        ("Documentation", "Swagger(OpenAPI)"),
        ("Test", "JUnit5, Playwright"),
        ("Version Control", "GitHub"),
        ("ERD", "ERDCloud"),
        ("UI 설계", "Figma")
    ]
    create_styled_table(doc, tech_headers, tech_data, [Mm(50), Mm(110)])
    
    # 5. 메뉴 구조
    add_heading_1(doc, "메뉴 구조")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    run = p.add_run("대시보드, 러닝 기록, 목표 관리, 대회 기록, 통계 분석, 마이페이지")
    set_font(run, "Malgun Gothic")
    
    # 6. 세부 메뉴 구성
    # 1. 대시보드
    add_heading_2(doc, "1. 대시보드")
    add_custom_bullet(doc, "로그인 후 첫 화면")
    add_custom_bullet(doc, "이번 달 목표")
    add_custom_bullet(doc, "월간 누적 거리")
    add_custom_bullet(doc, "월간 목표 달성률")
    add_custom_bullet(doc, "최근 운동 기록")
    add_custom_bullet(doc, "다가오는 레이스 목표")
    add_custom_bullet(doc, "현재 PB 현황")
    
    # 2. 러닝 기록
    add_heading_2(doc, "2. 러닝 기록")
    add_custom_bullet(doc, "기록 등록", num_label="•")
    add_custom_bullet(doc, "날짜", level=1)
    add_custom_bullet(doc, "운동시간", level=1)
    add_custom_bullet(doc, "거리", level=1)
    add_custom_bullet(doc, "평균페이스", level=1)
    add_custom_bullet(doc, "평균심박", level=1)
    add_custom_bullet(doc, "훈련종류 (회복주, 이지런, LSD, 템포런, 인터벌, 빌드업런, 레이스, 기타)", level=1)
    add_custom_bullet(doc, "훈련강도(RPE) (RPE 1 ~ 10)", level=1)
    add_custom_bullet(doc, "메모", level=1)
    add_custom_bullet(doc, "운동환경 (온도, 습도, 날씨 상태 / 기상 API 연동 검토)", level=1)
    
    add_custom_bullet(doc, "기록 조회", num_label="•")
    add_custom_bullet(doc, "일별: 기록 목록 조회", level=1)
    add_custom_bullet(doc, "주간: 주간 거리, 주간 운동 횟수, 주간 평균 페이스", level=1)
    add_custom_bullet(doc, "월간: 월 누적 거리, 월 평균 페이스, 월 평균 심박", level=1)
    
    # 3. 목표 관리
    add_heading_2(doc, "3. 목표 관리")
    add_custom_bullet(doc, "월간 목표", num_label="•")
    add_custom_bullet(doc, "연/월 기준 목표 설정", level=1)
    add_custom_bullet(doc, "목표 거리", level=1)
    add_custom_bullet(doc, "목표 운동 횟수", level=1)
    add_custom_bullet(doc, "예시: 2026년 7월 / 목표 거리 200km / 목표 운동 횟수 20회", level=1)
    
    add_custom_bullet(doc, "레이스 목표", num_label="•")
    add_custom_bullet(doc, "5km, 10km, 하프, 풀코스", level=1)
    add_custom_bullet(doc, "예시: 10km 목표 50분 / 하프 목표 1시간 50분 / 풀 목표 3시간 59분", level=1)
    
    # 4. 대회 기록 관리
    add_heading_2(doc, "4. 대회 기록 관리")
    add_custom_bullet(doc, "대회명")
    add_custom_bullet(doc, "대회일")
    add_custom_bullet(doc, "거리")
    add_custom_bullet(doc, "기록")
    add_custom_bullet(doc, "기록증 이미지")
    add_custom_bullet(doc, "메모")
    add_custom_bullet(doc, "추가 검토: 대회 장소, 대회 URL")
    add_custom_bullet(doc, "예시: JTBC 마라톤, 42.195km, 4:03:12")
    
    # 5. 통계 분석
    add_heading_2(doc, "5. 통계 분석")
    add_custom_bullet(doc, "개인 베스트(PB): 5km PB, 10km PB, 하프 PB, 풀코스 PB")
    add_custom_bullet(doc, "훈련 유형 분석: 훈련 유형 분포, LSD 비율, 인터벌 비율, 템포런 비율")
    add_custom_bullet(doc, "거리 분석: 주간 거리 추이, 월간 거리 추이, 연간 거리 추이")
    add_custom_bullet(doc, "목표 분석: 목표 달성률, 월별 목표 달성 추이")
    add_custom_bullet(doc, "페이스 분석: 평균 페이스 변화, 거리별 페이스 변화")
    add_custom_bullet(doc, "비교 분석: 전월 비교, 전년 동월 비교 (예시: 거리, 평균 페이스, 평균 심박)")
    add_custom_bullet(doc, "환경 분석: 온도별 페이스, 습도별 페이스, 날씨별 기록 비교")
    add_custom_bullet(doc, "러닝 레벨 분석 (재미 요소): 동물 캐릭터 기반 레벨 (예시: 고양이 러너, 사슴 러너, 늑대 러너, 치타 러너)")
    
    # 6. 마이페이지
    add_heading_2(doc, "6. 마이페이지")
    add_custom_bullet(doc, "프로필 조회")
    add_custom_bullet(doc, "회원 정보 수정")
    add_custom_bullet(doc, "소셜 계정 확인")
    
    # 7. 향후 확장
    add_heading_1(doc, "향후 확장 계획")
    add_heading_2(doc, "향후 확장(V2)")
    add_custom_bullet(doc, "자동 데이터 수집: 날씨 자동 수집 (기상 API 연동)")
    add_custom_bullet(doc, "FIT 파일 업로드: Garmin FIT 파일 업로드 (훈련 자동 등록)")
    add_custom_bullet(doc, "랩 분석: 1km 랩, 구간별 심박, 구간별 페이스 (예시: 1km 5:30, 2km 5:20, 3km 5:15)")
    
    add_heading_2(doc, "향후 확장(V3)")
    add_custom_bullet(doc, "AI 러닝 코치: 훈련 패턴 분석, 과훈련 감지, 회복 추천, 레이스 목표 달성 가능성 분석 (예시: 최근 4주 평균 거리, 최근 훈련 강도, 목표 기록 기반으로 Sub4 가능성 분석)")
    
    # 8. 예상 핵심 테이블
    add_heading_1(doc, "예상 핵심 테이블")
    add_custom_bullet(doc, "핵심 테이블: USER, RUN_RECORD, GOAL, RACE_RECORD")
    add_custom_bullet(doc, "향후 추가 테이블 (FIT 파일 연동 시): RUN_LAP")
    
    add_heading_2(doc, "테이블 구조 및 관계")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.line_spacing = 1.25
    run1 = p.add_run("USER\n")
    set_font(run1, "Malgun Gothic")
    run1.font.bold = True
    
    run2 = p.add_run(" ├─ RUN_RECORD\n ├─ GOAL\n └─ RACE_RECORD\n\nRUN_RECORD\n └─ RUN_LAP")
    set_font(run2, "Malgun Gothic")
    
    # 9. 우선 작성할 문서
    add_heading_1(doc, "우선 작성할 문서")
    docs = [
        "01. 프로젝트 개요서",
        "02. 요구사항 명세서",
        "03. ERD",
        "04. 테이블 명세서",
        "05. API 명세서",
        "06. 테스트 전략서"
    ]
    for d in docs:
        add_custom_bullet(doc, d)
        
    # 10. 개발 순서
    add_heading_1(doc, "개발 순서")
    steps = [
        "프로젝트 개요서 작성",
        "요구사항 명세서 작성",
        "ERD 설계",
        "테이블 명세서 작성",
        "API 설계",
        "Spring Boot 개발",
        "React 개발",
        "테스트 코드 작성",
        "통계 기능 구현",
        "FIT 업로드 확장"
    ]
    for idx, s in enumerate(steps, 1):
        add_custom_bullet(doc, s, num_label=f"{idx}.")
        
    # Document Footer / Info
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(24)
    run_footer = p_footer.add_run("Running Coach 기획서")
    set_font(run_footer, "Malgun Gothic")
    run_footer.font.size = Pt(8.5)
    run_footer.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    filename = "프로젝트_상세_설계_및_기획서_Running_Coach.docx"
    doc.save(filename)
    print(f"Successfully created strict document '{filename}'")

if __name__ == "__main__":
    main()
