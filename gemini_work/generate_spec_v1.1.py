import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helper: Set font for runs including East Asian (Korean) fonts
def set_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# Helper: Add custom bottom border to paragraphs for titles
def add_bottom_border(paragraph, color_hex="1A365D", size="12"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)  # 12 = 1.5 pt
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

# Helper: Set cell shading background color
def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

# Helper: Set custom padding for table cells
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper: Set custom cell borders
def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style is not None:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), border_style.get('val', 'single'))
            border_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(border_el)
        else:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
    tcPr.append(tcBorders)

# Helper: Set column widths for tables
def set_col_widths(table, widths):
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

# Helper: Add Heading 1 with Bottom Line
def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    add_bottom_border(p, color_hex="1A365D", size="10")
    return p

# Helper: Add Heading 2
def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    return p

# Helper: Add a custom list/bullet point
def add_custom_bullet(doc, lead_in, description, num_label=None, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6 * (level + 1))
    p.paragraph_format.first_line_indent = Mm(-6)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    
    if num_label:
        prefix = f"{num_label}  "
        color = RGBColor(0x1A, 0x36, 0x5D)
    else:
        prefix = "•  " if level == 0 else "○  "
        color = RGBColor(0x2B, 0x6C, 0xB0) if level == 0 else RGBColor(0x4A, 0x55, 0x68)
        
    prefix_run = p.add_run(prefix)
    prefix_run.font.bold = True
    prefix_run.font.color.rgb = color
    set_font(prefix_run, "Arial" if not num_label else "Malgun Gothic")
    
    if lead_in:
        lead_run = p.add_run(lead_in)
        set_font(lead_run, "Malgun Gothic")
        lead_run.font.bold = True
        lead_run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
    desc_run = p.add_run(description)
    set_font(desc_run, "Malgun Gothic")
    desc_run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

# Helper: Add page number field to word header/footer
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_table_from_data(doc, headers, data, col_widths, align_cols=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, "1A365D")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        set_cell_borders(cell, 
                         top={'val': 'single', 'sz': 4, 'color': '1A365D'},
                         bottom={'val': 'single', 'sz': 8, 'color': '1A365D'},
                         left=None, right=None)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        set_font(run, "Malgun Gothic")
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Style Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, text in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_borders(cell, 
                             top=None,
                             bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             left=None, right=None)
            
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(0)
            
            # Alignments
            if align_cols and c_idx < len(align_cols):
                p.alignment = align_cols[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
                
            run = p.add_run(text)
            set_font(run, "Malgun Gothic")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
                
    set_col_widths(table, col_widths)
    
    # Bottom margin spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)

def main():
    doc = Document()
    
    # Page setup - A4 with 25mm margins
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        
        # Enable different cover page header/footer
        section.different_first_page_header_footer = True
        
        # Configure Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hrun = hp.add_run("Running Coach — 프로젝트 산출물 문서 v1.1")
        set_font(hrun, "Malgun Gothic")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        
        # Configure Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run()
        add_page_number(frun)
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
        
    # Default Paragraph Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    
    cover_p1 = doc.add_paragraph()
    cover_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_p1.paragraph_format.space_after = Pt(6)
    r1 = cover_p1.add_run("프로젝트 산출물 문서")
    set_font(r1, "Malgun Gothic")
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_after = Pt(12)
    r2 = cover_title.add_run("Running Coach")
    set_font(r2, "Malgun Gothic")
    r2.font.size = Pt(28)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    cover_sub = doc.add_paragraph()
    cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_sub.paragraph_format.space_after = Pt(24)
    r3 = cover_sub.add_run("소셜 로그인 · 러닝 기록 및 목표 관리 · 대회 아카이빙 · 훈련 통계 분석")
    set_font(r3, "Malgun Gothic")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # Quick Navigation Bar Table
    nav_table = doc.add_table(rows=1, cols=4)
    nav_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    nav_widths = [Mm(40), Mm(40), Mm(40), Mm(40)]
    for idx, cell in enumerate(nav_table.rows[0].cells):
        set_cell_shading(cell, "1A365D")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell, top={'val': 'single', 'sz': 4, 'color': '1A365D'},
                               bottom={'val': 'single', 'sz': 4, 'color': '1A365D'},
                               left=None, right=None)
        cell.width = nav_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        nav_titles = ["요구사항 분석", "요구사항 명세서", "ER 다이어그램", "테스트 전략서"]
        r = p.add_run(nav_titles[idx])
        set_font(r, "Malgun Gothic")
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    
    # Cover Metadata Table
    meta_data = [
        ("프로젝트명", "Running Coach (러닝 코치)"),
        ("기술 스택", "Spring Boot 3.x / Spring Security / JWT / MyBatis / React / PostgreSQL / OAuth2 / Playwright"),
        ("개발 환경", "Java 17, Node.js 24.x, IntelliJ IDEA, VS Code, Git/GitHub"),
        ("문서 버전", "v1.1"),
        ("작성 기준", "산출물 가이드(산출물가이드.pdf) v1.0 규격 기준 기획서 개정 작성")
    ]
    
    m_table = doc.add_table(rows=len(meta_data), cols=2)
    m_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (lbl, val) in enumerate(meta_data):
        row_cells = m_table.rows[r_idx].cells
        
        # Label cell
        cell_lbl = row_cells[0]
        set_cell_shading(cell_lbl, "F0F4F8")
        set_cell_margins(cell_lbl, top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell_lbl, top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                                   bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                                   left=None, right=None)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_after = Pt(0)
        r_lbl = p_lbl.add_run(lbl)
        set_font(r_lbl, "Malgun Gothic")
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        # Value cell
        cell_val = row_cells[1]
        set_cell_margins(cell_val, top=100, bottom=100, left=150, right=150)
        set_cell_borders(cell_val, top={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                                   bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                                   left=None, right=None)
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_val.paragraph_format.line_spacing = 1.2
        p_val.paragraph_format.space_after = Pt(0)
        r_val = p_val.add_run(val)
        set_font(r_val, "Malgun Gothic")
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
        
    set_col_widths(m_table, [Mm(40), Mm(120)])
    
    # Page Break to start Section 1
    doc.add_page_break()
    
    # ----------------------------------------------------
    # 1. 요구사항 분석 (Requirements Analysis)
    # ----------------------------------------------------
    add_heading_1(doc, "1. 요구사항 분석 (Requirements Analysis)")
    
    # 1.1 프로젝트 개요
    add_heading_2(doc, "1.1 프로젝트 개요")
    p_overview = doc.add_paragraph()
    p_overview.paragraph_format.left_indent = Mm(6)
    p_overview.paragraph_format.line_spacing = 1.3
    p_overview.paragraph_format.space_after = Pt(8)
    r_overview = p_overview.add_run(
        "본 프로젝트는 Spring Boot 기반으로 소셜 로그인(OAuth2), 러닝 기록 등록/수정/삭제, 월별 및 레이스 목표 관리, "
        "대회 기록 및 기록증 이미지 아카이빙, 그리고 훈련 성과 통계 분석 기능을 제공하는 러닝 훈련 관리 플랫폼이다. "
        "MyBatis 데이터 매퍼를 활용해 PostgreSQL 및 Oracle 데이터베이스와의 다중 DB 연동 안정성을 확보하며, "
        "프론트엔드 React 구조와 통합하여 사용자에게 고가독성 훈련 분석 통계 및 직관적인 대시보드 환경을 전달하는 것을 목표로 한다."
    )
    set_font(r_overview, "Malgun Gothic")
    r_overview.font.size = Pt(10)
    
    # 1.2 추진 배경 및 목적
    add_heading_2(doc, "1.2 추진 배경 및 목적")
    add_custom_bullet(doc, "", "기존의 상용 러닝 앱들이 단순한 달리기 활동 기록(거리, 시간) 저장 중심으로 설계되어 있어, 월간 목표 수립이나 목표 대비 달성률 관리, 참가 예정 대회 타겟 기록 설계 등 고도화된 훈련 관리 피드백을 원하는 러너들의 니즈를 해소하기에 한계가 있음.")
    add_custom_bullet(doc, "", "러너가 매월 설정하는 목표 거리 및 훈련 횟수와 레이스 타겟 시간(5k, 10k, 하프, 풀코스 등)에 대한 페이스 차트를 설계하고 이를 통계적으로 분석해줄 수 있는 전용 분석 아카이빙 플랫폼의 필요성 증대.")
    add_custom_bullet(doc, "", "대회 참가 이력과 공식 기록증 이미지를 하나의 보관 공간에서 체계적으로 디지털 이력화하여 개인 러닝 성장의 타임라인을 제공하고자 함.")
    add_custom_bullet(doc, "", "기본 로그인 및 단편적인 훈련 기록 관리 단계에 머무르지 않고, 차트와 환경 통계 분석(온도, 습도, 기상 요인 등)을 유기적으로 연동하며 향후 가민 스마트 워치(FIT) 연동 및 AI 코칭으로 유연하게 연계할 수 있는 백엔드/프론트엔드 표준 아키텍처 구축을 목적으로 함.")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # 1.3 사용자(이해관계자) 분석
    add_heading_2(doc, "1.3 사용자(이해관계자) 분석")
    user_headers = ["구분", "설명", "주요 니즈"]
    user_data = [
        ("비로그인 사용자", "소셜 로그인을 연동하기 전 상태의 일반 방문자", "서비스의 기획 목표 탐색 및 구글/네이버를 활용한 간편 소셜 로그인 가입 접근"),
        ("일반 회원 (러너)", "로그인을 완료한 모든 실사용자(러너)", "개인 훈련 기록 CRUD, 연/월간 정량적 주행 목표 설정 및 실시간 달성률 분석 차트 조회"),
        ("대회 참가 러너", "공식 마라톤 대회 참가를 병행하는 러너", "대회 스케줄 D-Day 조회, 참가 완주 기록 및 실물 기록증 아카이빙, 개인 최고기록(PB) 관리"),
        ("시스템 관리자", "플랫폼의 데이터 모니터링 및 운영 주체", "회원 통계 분석 모니터링, 외부 기상 API 연동 상태 및 시스템 백업 인프라 관리")
    ]
    create_table_from_data(doc, user_headers, user_data, [Mm(30), Mm(65), Mm(65)], 
                           [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    
    # 1.4 As-Is / To-Be 분석
    add_heading_2(doc, "1.4 As-Is / To-Be 분석")
    asis_headers = ["구분", "As-Is (기록 부재 및 단순 텍스트 저장)", "To-Be (본 시스템 - Running Coach v1.1)"]
    asis_data = [
        ("훈련 기록 관리", "단순 텍스트 중심의 메모나 파편화된 기기로 날짜와 달린 거리만 파편화하여 기록", "날짜, 시간, 거리, 평균 페이스, 평균 심박, 훈련 종류 및 RPE(주관적 훈련 강도) 통합 관리"),
        ("목표 진척도 분석", "목표 거리를 머릿속으로만 계산하거나 별도의 통계적인 피드백을 받지 못함", "연/월 기준 목표 거리/횟수 설정 및 실시간 누적 거리 대비 진척도 시각화 대시보드 제공"),
        ("대회 이력 보관", "완주 기록과 실물 기록증을 하드카피로 보관하거나 SNS 일회성 업로드에 그침", "대회명, 일시, 거리, 기록과 함께 공식 디지털 기록증 이미지를 웹 아카이브에 영구 아카이빙"),
        ("기록 데이터 분석", "달린 기록이 날짜순 리스트로만 나열되어 패턴 파악이 불가", "5k/10k/Half/Full 거리별 PB 자동 갱신 및 주/월/연간 거리 추이, 온도/날씨별 상관 분석 제공"),
        ("게이미케이션 요소", "동기부여 수단이 부재하여 자발적 훈련 유지가 어려움", "누적 주행 거리에 따른 동물 캐릭터 기반 러닝 레벨 시스템(고양이~치타) 도입으로 흥미 유발")
    ]
    create_table_from_data(doc, asis_headers, asis_data, [Mm(30), Mm(65), Mm(65)],
                           [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    
    # 1.5 핵심 기능 요구사항 요약
    add_heading_2(doc, "1.5 핵심 기능 요구사항 요약")
    reqs = [
        ("FR-01. 소셜 로그인 및 회원 인증: ", "네이버 및 구글 계정을 연동한 간편 소셜 로그인 및 소셜 회원 정보 연동 기능"),
        ("FR-02. 러닝 훈련 기록 등록: ", "날짜, 시간, 거리, 페이스, 심박, 메모, 훈련종류(회복주~레이스), 훈련강도(RPE 1~10) 및 당일 환경(온도, 습도, 기상상태) 수동/자동 입력"),
        ("FR-03. 주간/월간 기록 집계 조회: ", "캘린더/일별 리스트 뷰 및 주간 누적 거리/횟수/페이스, 월간 누적 주행 통계 조회"),
        ("FR-04. 정량 목표 설정 및 관리: ", "연/월별 타겟 거리(km) 및 훈련 횟수 설정 기능"),
        ("FR-05. 레이스 타겟 페이스 계획: ", "5k, 10k, 하프, 풀코스 완주 목표 시간 설정 및 타겟 페이스 산출 기능"),
        ("FR-06. 마라톤 대회 정보 아카이빙: ", "대회명, 날짜, 거리, 완주기록 등록, 대회 장소 및 공식 홈페이지 URL 연동, 실물 기록증 이미지 업로드 관리"),
        ("FR-07. 개인 최고 기록(PB) 관리: ", "훈련 및 대회 데이터에서 거리별(5km, 10km, Half, Full) 최고 기록 자동 추출 및 대시보드 상시 표출"),
        ("FR-08. 다차원 통계 시각화 및 레벨링: ", "주/월/년 거리 변화 그래프, 전월/전년 동월 비교 분석, 날씨 조건별 페이스 영향도 분석 및 캐릭터 러닝 레벨(고양이~치타) 갱신")
    ]
    for r_lbl, r_desc in reqs:
        add_custom_bullet(doc, r_lbl, r_desc)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # 1.6 제약사항 (Constraints)
    add_heading_2(doc, "1.6 제약사항 (Constraints)")
    add_custom_bullet(doc, "데이터베이스: ", "MySQL 8 및 Oracle DB 다중 데이터베이스 환경 지원 및 표준 SQL DDL/DML 호환성 준수 (MyBatis 연동)")
    add_custom_bullet(doc, "파일 저장 방식: ", "사용자가 업로드하는 대회 기록증 이미지는 지정된 로컬 파일시스템 경로(예: C:/uploads/certificates/)에 저장하며, 클라우드 스토리지(S3 등)는 v1.1 개발 범위에서 배제함")
    add_custom_bullet(doc, "인증 및 권한 통제: ", "일반 자체 회원가입 기능을 생략하고 소셜 OAuth2 인증 프로토콜을 단일 진입점으로 설정하며, Spring Security로 보호함. 타 사용자의 비공개 레코드 및 기록증 접근을 서버 단에서 강제 통제")
    add_custom_bullet(doc, "파일 업로드 한계치 설정: ", "1개 대회 기록증 파일 크기는 최대 10MB로 제한하고, 멀티 파트 요청 전체 최대 크기는 20MB로 제한함 (application.yml 환경설정 기준)")
    add_custom_bullet(doc, "주요 엔티티 기본키(PK) 전략: ", "각 테이블의 기본키 생성은 DBMS의 자동 증가 기능을 활용하는 IDENTITY 전략(MySQL AUTO_INCREMENT / Oracle SEQUENCE 또는 IDENTITY)을 사용함")
    
    # Document Footer / Info
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(36)
    run_footer = p_footer.add_run("문서 작성일: 2026. 06. 22  |  작성자: Running Coach 기획팀")
    set_font(run_footer, "Malgun Gothic")
    run_footer.font.size = Pt(8.5)
    run_footer.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    filename = "프로젝트_개요서_Running_Coach_v1.1.docx"
    doc.save(filename)
    print(f"Successfully created revised document '{filename}'")

if __name__ == "__main__":
    main()
