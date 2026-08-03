import os
from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Helper: Set font for runs including East Asian (Korean) fonts
def set_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# Helper: Add custom bottom border to paragraphs for titles
def add_bottom_border(paragraph, color_hex="1A365D", size="12"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)  # 12 = 1.5 pt
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

# Helper: Set cell shading background color
def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

# Helper: Set custom padding for table cells
def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper: Set custom cell borders
def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style is not None:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), border_style.get('val', 'single'))
            border_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            border_el.set(qn('w:space'), '0')
            border_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(border_el)
        else:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
    tcPr.append(tcBorders)

# Helper: Set column widths for tables
def set_col_widths(table, widths):
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w

# Helper: Add Heading 1 with Bottom Line
def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    add_bottom_border(p, color_hex="1A365D", size="8")
    return p

# Helper: Add Heading 2
def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    set_font(run, "Malgun Gothic")
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
    return p

# Helper: Add a custom list/bullet point
def add_custom_bullet(doc, text, num_label=None, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6 * (level + 1))
    p.paragraph_format.first_line_indent = Mm(-6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    
    if num_label:
        prefix = f"{num_label}  "
        color = RGBColor(0x1A, 0x36, 0x5D)
    else:
        prefix = "-  "
        color = RGBColor(0x2B, 0x6C, 0xB0)
        
    prefix_run = p.add_run(prefix)
    prefix_run.font.bold = True
    prefix_run.font.color.rgb = color
    set_font(prefix_run, "Arial" if not num_label else "Malgun Gothic")
    
    desc_run = p.add_run(text)
    set_font(desc_run, "Malgun Gothic")
    desc_run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

# Helper: Add page number field to word header/footer
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

# Helper: Create standard styled DB table specification grid
def create_db_table_grid(doc, headers, data, col_widths):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_shading(cell, "1A365D")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        set_cell_borders(cell, 
                         top={'val': 'single', 'sz': 4, 'color': '1A365D'},
                         bottom={'val': 'single', 'sz': 8, 'color': '1A365D'},
                         left=None, right=None)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        set_font(run, "Malgun Gothic")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, text in enumerate(row_data):
            cell = row_cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            set_cell_borders(cell, 
                             top=None,
                             bottom={'val': 'single', 'sz': 4, 'color': 'E2E8F0'},
                             left=None, right=None)
            
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            
            if c_idx in (0, 2, 3, 4):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            run = p.add_run(text)
            set_font(run, "Malgun Gothic")
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
            
            # Highlight key fields
            if c_idx == 0:
                run.font.bold = True
            if c_idx == 4 and text: # Key column (PK/FK)
                run.font.bold = True
                if text == "PK":
                    run.font.color.rgb = RGBColor(0x9B, 0x2C, 0x2C)
                elif text == "FK":
                    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
                    
    set_col_widths(table, col_widths)
    
    # Bottom margin spacing
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(10)

def main():
    doc = Document()
    
    # Page setup - A4 with 25mm margins
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        
        # Configure Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hrun = hp.add_run("Running Coach — 데이터베이스 테이블 명세서 v1.1")
        set_font(hrun, "Malgun Gothic")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        
        # Configure Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run()
        add_page_number(frun)
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
        
    # Default Paragraph Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Malgun Gothic'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(15)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("데이터베이스 테이블 명세서")
    set_font(title_run, "Malgun Gothic")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(20)
    sub_run = subtitle_p.add_run("Running Coach 데이터 모델 설계 산출물 (v1.1 개정안)")
    set_font(sub_run, "Malgun Gothic")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    add_bottom_border(subtitle_p, color_hex="D2D6DC", size="6")
    
    # ----------------------------------------------------
    # Section 1. 테이블 정의 (Table Definitions)
    # ----------------------------------------------------
    add_heading_1(doc, "1. 테이블 정의")
    
    grid_headers = ["Field", "Field2 (한글명)", "Domain", "Type", "Key"]
    grid_widths = [Mm(35), Mm(35), Mm(30), Mm(40), Mm(20)]
    
    # RUNNER
    add_heading_2(doc, "RUNNER (러너 정보)")
    runner_data = [
        ("runner_id", "러너ID", "PK", "BIGINT", "PK"),
        ("nickname", "닉네임", "NAME", "VARCHAR(50)", ""),
        ("profile_image_url", "프로필이미지URL", "URL", "VARCHAR(500)", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, runner_data, grid_widths)
    
    # RUNNER_SOCIAL_ACCOUNT
    add_heading_2(doc, "RUNNER_SOCIAL_ACCOUNT (소셜 로그인 정보)")
    social_data = [
        ("social_account_id", "소셜계정ID", "PK", "BIGINT", "PK"),
        ("runner_id", "러너ID", "FK", "BIGINT", "FK"),
        ("provider", "제공자", "CODE", "VARCHAR(20)", ""),
        ("provider_user_id", "제공자사용자ID", "ID", "VARCHAR(100)", ""),
        ("provider_email", "소셜이메일", "EMAIL", "VARCHAR(200)", ""),
        ("is_primary", "대표계정여부", "YN", "CHAR(1)", ""),
        ("connected_at", "연동일시", "DATETIME", "TIMESTAMP", ""),
        ("last_login_at", "최종로그인일시", "DATETIME", "TIMESTAMP", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, social_data, grid_widths)
    
    # RUN_RECORD
    add_heading_2(doc, "RUN_RECORD (러닝 기록)")
    run_record_data = [
        ("run_record_id", "러닝기록ID", "PK", "BIGINT", "PK"),
        ("runner_id", "러너ID", "FK", "BIGINT", "FK"),
        ("run_datetime", "운동일시", "DATETIME", "TIMESTAMP", ""),
        ("run_date", "운동일자", "DATE", "DATE", ""),
        ("duration_sec", "운동시간(초)", "DURATION", "INTEGER", ""),
        ("distance_km", "거리(km)", "DISTANCE", "NUMERIC(6,2)", ""),
        ("avg_pace_sec", "평균페이스(초)", "PACE", "INTEGER", ""),
        ("avg_hr", "평균심박수", "HEART_RATE", "INTEGER", ""),
        ("training_type_code", "훈련유형코드", "CODE", "VARCHAR(30)", ""),
        ("rpe", "운동강도", "RPE", "INTEGER", ""),
        ("temperature", "온도", "TEMPERATURE", "NUMERIC(4,1)", ""),
        ("humidity", "습도", "HUMIDITY", "INTEGER", ""),
        ("weather_code", "날씨코드", "CODE", "VARCHAR(30)", ""),
        ("memo", "메모", "TEXT", "VARCHAR(1000)", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, run_record_data, grid_widths)
    
    # GOAL
    add_heading_2(doc, "GOAL (목표 관리)")
    goal_data = [
        ("goal_id", "목표ID", "PK", "BIGINT", "PK"),
        ("runner_id", "러너ID", "FK", "BIGINT", "FK"),
        ("goal_type_code", "목표유형코드", "CODE", "VARCHAR(30)", ""),
        ("goal_unit_code", "목표단위코드", "CODE", "VARCHAR(30)", ""),
        ("target_year", "목표년도", "YEAR", "INTEGER", ""),
        ("target_month", "목표월", "MONTH", "INTEGER", ""),
        ("target_value", "목표값", "TARGET", "NUMERIC(10,2)", ""),
        ("achieved_value", "달성값", "RESULT", "NUMERIC(10,2)", ""),
        ("start_date", "시작일", "DATE", "DATE", ""),
        ("end_date", "종료일", "DATE", "DATE", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, goal_data, grid_widths)
    
    # RACE_RECORD
    add_heading_2(doc, "RACE_RECORD (대회 기록)")
    race_data = [
        ("race_record_id", "대회기록ID", "PK", "BIGINT", "PK"),
        ("runner_id", "러너ID", "FK", "BIGINT", "FK"),
        ("race_name", "대회명", "NAME", "VARCHAR(200)", ""),
        ("race_date", "대회일자", "DATE", "DATE", ""),
        ("race_type_code", "대회유형코드", "CODE", "VARCHAR(30)", ""),
        ("record_sec", "기록(초)", "RECORD", "INTEGER", ""),
        ("certificate_image_url", "기록증이미지URL", "URL", "VARCHAR(500)", ""),
        ("memo", "메모", "TEXT", "VARCHAR(1000)", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, race_data, grid_widths)
    
    # CODE_GROUP
    add_heading_2(doc, "CODE_GROUP (코드 그룹)")
    code_group_data = [
        ("group_code", "그룹코드", "PK", "VARCHAR(30)", "PK"),
        ("group_name", "그룹명", "NAME", "VARCHAR(100)", ""),
        ("description", "설명", "TEXT", "VARCHAR(500)", ""),
        ("use_yn", "사용여부", "YN", "CHAR(1)", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, code_group_data, grid_widths)
    
    # CODE_DETAIL
    add_heading_2(doc, "CODE_DETAIL (코드 상세)")
    code_detail_data = [
        ("code_id", "코드ID", "PK", "BIGINT", "PK"),
        ("group_code", "그룹코드", "FK", "VARCHAR(30)", "FK"),
        ("code_value", "코드값", "CODE", "VARCHAR(30)", ""),
        ("code_name", "코드명", "NAME", "VARCHAR(100)", ""),
        ("description", "설명", "TEXT", "VARCHAR(500)", ""),
        ("sort_order", "정렬순서", "ORDER", "INTEGER", ""),
        ("use_yn", "사용여부", "YN", "CHAR(1)", ""),
        ("created_at", "생성일시", "DATETIME", "TIMESTAMP", ""),
        ("updated_at", "수정일시", "DATETIME", "TIMESTAMP", "")
    ]
    create_db_table_grid(doc, grid_headers, code_detail_data, grid_widths)

    # ----------------------------------------------------
    # Section 2. 관계도 (Relationship Model)
    # ----------------------------------------------------
    add_heading_1(doc, "2. 관계도 (현재 기준)")
    
    p_erd = doc.add_paragraph()
    p_erd.paragraph_format.left_indent = Mm(6)
    p_erd.paragraph_format.line_spacing = 1.3
    
    r_erd = p_erd.add_run(
        "■ 테이블 간 계층 관계\n"
        "RUNNER (러너 정보)\n"
        " ├── RUNNER_SOCIAL_ACCOUNT (소셜 로그인 정보)\n"
        " ├── RUN_RECORD (러닝 기록)\n"
        " ├── GOAL (목표 관리)\n"
        " └── RACE_RECORD (대회 기록)\n\n"
        "CODE_GROUP (코드 그룹)\n"
        " └── CODE_DETAIL (코드 상세)\n\n"
        "■ 공통 코드 매핑 관계\n"
        "RUN_RECORD (러닝 기록)\n"
        " ├── training_type_code → CODE_DETAIL\n"
        " └── weather_code → CODE_DETAIL\n\n"
        "GOAL (목표 관리)\n"
        " ├── goal_type_code → CODE_DETAIL\n"
        " └── goal_unit_code → CODE_DETAIL\n\n"
        "RACE_RECORD (대회 기록)\n"
        " └── race_type_code → CODE_DETAIL\n"
    )
    set_font(r_erd, "Consolas")
    r_erd.font.size = Pt(9.5)
    r_erd.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)

    # ----------------------------------------------------
    # Section 3. 공통 코드 데이터 정의 (Common Code Data)
    # ----------------------------------------------------
    add_heading_1(doc, "3. 공통 코드 데이터 정의")
    
    # Code Group
    add_heading_2(doc, "CODE_GROUP 초기 데이터")
    for group in ["TRAINING_TYPE: 훈련 유형 코드 그룹", "WEATHER_TYPE: 날씨 유형 코드 그룹", "GOAL_TYPE: 목표 유형 코드 그룹", "RACE_TYPE: 대회 유형 코드 그룹", "RUNNING_LEVEL: 러닝 레벨 코드 그룹"]:
        add_custom_bullet(doc, group)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    # Code Details: GOAL_TYPE
    add_heading_2(doc, "GOAL_TYPE 상세 코드")
    goals = [
        "MONTH_DISTANCE: 월간 목표 거리",
        "MONTH_RUN_COUNT: 월간 목표 운동 횟수",
        "RACE_5K: 5km 목표 기록",
        "RACE_10K: 10km 목표 기록",
        "RACE_HALF: 하프 마라톤 목표 기록",
        "RACE_FULL: 풀코스 마라톤 목표 기록"
    ]
    for g in goals:
        add_custom_bullet(doc, g)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    # Code Details: TRAINING_TYPE
    add_heading_2(doc, "TRAINING_TYPE 상세 코드")
    trainings = [
        "RECOVERY: 회복주 (피로 회복 목적 저강도 러닝)",
        "EASY: 이지런 (가벼운 조깅)",
        "LSD: Long Slow Distance (장거리 지구력 훈련)",
        "TEMPO: 템포런 (페이스 유지 지속주)",
        "INTERVAL: 인터벌 (고강도 반복 훈련)",
        "BUILDUP: 빌드업런 (점진적 페이스 가속 훈련)",
        "RACE: 레이스 (대회 참가 주행)",
        "OTHER: 기타"
    ]
    for t in trainings:
        add_custom_bullet(doc, t)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
        
    # Code Details: WEATHER_TYPE
    add_heading_2(doc, "WEATHER_TYPE 상세 코드")
    weathers = [
        "SUNNY: 맑음",
        "CLOUDY: 흐림",
        "RAIN: 비",
        "SNOW: 눈"
    ]
    for w in weathers:
        add_custom_bullet(doc, w)
        
    # Document Footer / Info
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_footer.paragraph_format.space_before = Pt(36)
    run_footer = p_footer.add_run("문서 작성일: 2026. 06. 23  |  작성자: Running Coach 기획팀")
    set_font(run_footer, "Malgun Gothic")
    run_footer.font.size = Pt(8.5)
    run_footer.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
    
    filename = "c:/Runners_plan/산출물/테이블_명세서_Running_Coach_v1.1.docx"
    doc.save(filename)
    print(f"Successfully created revised Table Specification document '{filename}'")

if __name__ == "__main__":
    main()
